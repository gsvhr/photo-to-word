"""Генератор Word документов с фотографиями."""

import os
import io
import tempfile
from typing import Optional, Callable
from pathlib import Path
import logging
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class WordGenerator:
    """Класс для генерации Word документов."""

    def __init__(self, config, image_processor):
        self.config = config
        self.image_processor = image_processor
        self.document = None
        self.cancel_requested = False

    def create_document(
        self, orientation: str = "portrait", table_width_cm: float = 16
    ):
        """Создать новый документ с настройками."""
        self.document = Document()

        # Настройка полей (все по 1 см)
        section = self.document.sections[0]
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)

        # Настройка ориентации
        if orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = (
                section.page_height,
                section.page_width,
            )

        # Сохраняем ширину таблицы для использования в generate
        self.table_width_cm = table_width_cm

        return self.document

    def _remove_paragraph_spacing(self, paragraph):
        """Убрать отступы после абзаца."""
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0

    def generate(
        self,
        orientation: str = "portrait",
        quality: int = 85,
        table_width_cm: float = 16,
        progress_callback: Optional[Callable] = None,
    ) -> Optional[io.BytesIO]:
        """Сгенерировать документ и вернуть его как BytesIO объект."""

        self.cancel_requested = False

        # === 1. ПРОВЕРКА СОЗДАНИЯ ДОКУМЕНТА ===
        try:
            self.create_document(orientation, table_width_cm)
        except Exception as e:
            logger.error(f"Ошибка при создании документа: {e}")
            raise RuntimeError(f"Не удалось создать документ: {e}") from e

        # === 2. ПРОВЕРКА ЧТО ДОКУМЕНТ СУЩЕСТВУЕТ ===
        if self.document is None:
            raise RuntimeError("Документ не был создан (self.document = None)")

        # === 3. ПРОВЕРКА НАЛИЧИЯ ИЗОБРАЖЕНИЙ ===
        image_paths = self.image_processor.get_image_paths()
        total = len(image_paths)

        if total == 0:
            raise ValueError("Нет изображений для генерации")

        # Получаем углы поворота для всех изображений
        rotations = [self.image_processor.get_rotation(i) for i in range(total)]

        # Всегда 2 колонки
        cols = 2
        target_width = self.config.get_image_width(orientation)

        # Рассчитываем ширину колонки
        col_width_cm = table_width_cm / cols
        col_width = Cm(col_width_cm)

        # === 4. СОЗДАНИЕ ТАБЛИЦЫ (с проверкой) ===
        try:
            # Создаем таблицу с фиксированной шириной
            table = self.document.add_table(rows=0, cols=cols)
            table.autofit = False
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Устанавливаем ширину для всех колонок
            for col in table.columns:
                col.width = col_width
        except Exception as e:
            logger.error(f"Ошибка при создании таблицы: {e}")
            raise RuntimeError(f"Не удалось создать таблицу: {e}") from e

        # Рассчитываем количество строк, которые понадобятся
        rows_needed = (total + cols - 1) // cols  # округление вверх

        # Создаем все строки заранее
        for _ in range(rows_needed):
            table.add_row()

        # Список временных файлов для удаления
        temp_files = []

        try:
            for i, img_path in enumerate(image_paths):
                # Проверка отмены
                if self.cancel_requested:
                    logger.info("Генерация отменена")
                    return None

                # Обработка изображения
                try:
                    img = self.image_processor.process_for_word(
                        img_path, target_width, quality, rotations[i]
                    )
                except Exception as e:
                    logger.error(f"Ошибка обработки изображения {img_path}: {e}")
                    raise RuntimeError(
                        f"Не удалось обработать изображение {Path(img_path).name}: {e}"
                    ) from e

                # Сохранение во временный файл
                try:
                    with tempfile.NamedTemporaryFile(
                        suffix=".jpg", delete=False
                    ) as tmp_file:
                        img.save(tmp_file.name, "JPEG", quality=quality)
                        tmp_path = tmp_file.name
                        temp_files.append(tmp_path)
                except Exception as e:
                    logger.error(f"Ошибка сохранения временного файла: {e}")
                    raise RuntimeError(
                        f"Не удалось сохранить временный файл: {e}"
                    ) from e

                # Получаем строку и ячейку для текущего изображения
                row = table.rows[i // cols]
                cell = row.cells[i % cols]

                # === 5. ДОБАВЛЕНИЕ ИЗОБРАЖЕНИЯ ===
                try:
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self._remove_paragraph_spacing(paragraph)

                    run = paragraph.add_run()
                    run.add_picture(tmp_path, width=col_width)
                except Exception as e:
                    logger.error(f"Ошибка вставки изображения в ячейку: {e}")
                    raise RuntimeError(f"Не удалось вставить изображение: {e}") from e

                # === 6. ДОБАВЛЕНИЕ ПОДПИСИ ===
                try:
                    caption_text = self.config.get_caption_text(
                        i + 1, Path(img_path).stem
                    )
                    caption_para = cell.add_paragraph()
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    self._remove_paragraph_spacing(caption_para)

                    caption_run = caption_para.add_run(caption_text)
                    caption_run.font.name = self.config.word_document_defaults[
                        "font_name"
                    ]
                    caption_run.font.size = Pt(
                        self.config.word_document_defaults["font_size"]
                    )

                    # Защита от разрыва (фото + подпись на одной странице)
                    cell.paragraphs[0].paragraph_format.keep_with_next = True
                except Exception as e:
                    logger.error(f"Ошибка добавления подписи: {e}")
                    # Не критично, продолжаем
                    logger.warning(f"Пропуск подписи для изображения {i+1}")

                # Прогресс
                if progress_callback:
                    try:
                        progress_callback(i + 1, total)
                    except Exception as e:
                        logger.warning(f"Ошибка в progress_callback: {e}")

            # === 7. СОХРАНЕНИЕ В ПАМЯТЬ ===
            try:
                docx_bytes = io.BytesIO()
                self.document.save(docx_bytes)
                docx_bytes.seek(0)
                logger.info(
                    f"Документ успешно сгенерирован: {len(docx_bytes.getvalue())} байт"
                )
                return docx_bytes
            except Exception as e:
                logger.error(f"Ошибка сохранения документа в память: {e}")
                raise RuntimeError(f"Не удалось сохранить документ: {e}") from e

        finally:
            # === 8. УДАЛЕНИЕ ВРЕМЕННЫХ ФАЙЛОВ ===
            for tmp_path in temp_files:
                try:
                    if os.path.exists(tmp_path):
                        os.unlink(tmp_path)
                        logger.debug(f"Удален временный файл: {tmp_path}")
                except Exception as e:
                    logger.warning(f"Не удалось удалить временный файл {tmp_path}: {e}")

    def cancel(self):
        """Отменить генерацию."""
        self.cancel_requested = True

    def save_to_file(
        self,
        filepath: str,
        orientation: str = "portrait",
        quality: int = 85,
        table_width_cm: float = 16,
        progress_callback: Optional[Callable] = None,
    ) -> bool:
        """Сохранить документ в файл."""
        try:
            docx_bytes = self.generate(
                orientation, quality, table_width_cm, progress_callback
            )
            if docx_bytes:
                with open(filepath, "wb") as f:
                    f.write(docx_bytes.getvalue())
                return True
        except Exception as e:
            logger.error(f"Ошибка сохранения файла: {e}")
            raise
        return False
